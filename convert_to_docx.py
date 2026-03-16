from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

INPUT  = r"c:\xampp\htdocs\PetCloud\mini.txt"
OUTPUT = r"c:\xampp\htdocs\PetCloud\PetCloud_MiniProject_Report.docx"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1)

# ── Default font ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def set_font(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if level == 0:
            run.font.size = Pt(16)
        elif level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)
    return p

def add_table_from_lines(doc, rows):
    """Build a Word table from a list of column lists."""
    if not rows:
        return
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph()

# ── Detect line type ──────────────────────────────────────────────────────────
def is_separator(line):
    return re.match(r'^[-=+]{5,}', line.strip())

def is_chapter_heading(line):
    return re.match(r'^CHAPTER\s+\d+\s*$', line.strip())

def is_section_heading(line):
    # e.g.  "4.2.1 USE CASE DIAGRAM"  or  "1.1 PROJECT OVERVIEW"
    return re.match(r'^\d+(\.\d+)*\s+[A-Z]', line.strip())

def is_page_heading(line):
    # All-caps lines that serve as section titles (ABSTRACT, INTRODUCTION, etc.)
    s = line.strip()
    return s and s == s.upper() and len(s) > 3 and not re.match(r'^[-=+|]', s) and not re.search(r'\d{4}', s)

def is_table_divider(line):
    return re.match(r'^[-+]+[-+]+', line.strip())

def parse_table_row(line):
    """Split a pipe-delimited row into cells."""
    line = line.strip().strip('|')
    return [c.strip() for c in line.split('|')]

# ── Parse table blocks ────────────────────────────────────────────────────────
def parse_pipe_table(lines, start):
    """Return (rows_list, end_index) for a pipe-delimited table starting at start."""
    rows = []
    i = start
    while i < len(lines):
        l = lines[i].strip()
        if '|' in l and not is_separator(l):
            rows.append(parse_table_row(l))
        elif is_table_divider(l):
            pass  # skip divider lines
        else:
            break
        i += 1
    return rows, i

# ── Main conversion ───────────────────────────────────────────────────────────
with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Strip \r
lines = [l.rstrip('\r\n') for l in lines]

i = 0
chapter_title_next = False  # Flag: next non-empty line after CHAPTER X is the chapter title

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip pure separator lines
    if is_separator(line) and stripped.replace('-','').replace('=','').replace('+','') == '':
        i += 1
        continue

    # Empty line → small spacer
    if stripped == '':
        i += 1
        continue

    # "CHAPTER N" → Heading 1
    if is_chapter_heading(stripped):
        chapter_title_next = True
        i += 1
        continue

    # Chapter title (line after CHAPTER N)
    if chapter_title_next and stripped:
        add_heading(doc, stripped, level=1)
        chapter_title_next = False
        i += 1
        continue

    # Numbered section headings e.g. "4.2.1 USE CASE DIAGRAM"
    m = re.match(r'^(\d+(\.\d+)*)\s+(.*)', stripped)
    if m:
        number = m.group(1)
        title  = m.group(3)
        depth  = number.count('.')
        level  = min(depth + 1, 4)
        add_heading(doc, f"{number} {title}", level=level)
        i += 1
        continue

    # All-caps page-title headings (ABSTRACT, DECLARATION, etc.)
    if is_page_heading(stripped) and len(stripped.split()) <= 6:
        add_heading(doc, stripped, level=1)
        i += 1
        continue

    # Table-separator dashes lines (---- table borders) — detect pipe table block
    if '|' in stripped and not is_separator(stripped):
        # Gather all consecutive pipe lines
        rows = []
        while i < len(lines):
            l = lines[i].strip()
            if '|' in l:
                if not re.match(r'^[-+]+$', l.replace('|','').replace('+','').replace('-','')):
                    rows.append(parse_table_row(l))
            elif re.match(r'^[-+]{3,}', l):
                pass  # divider
            else:
                break
            i += 1
        if rows:
            add_table_from_lines(doc, rows)
        continue

    # Table design header "--------" separator with table name above it
    if re.match(r'^-{5,}$', stripped):
        i += 1
        continue

    # Regular paragraph / bullet
    para = doc.add_paragraph()
    # Detect bullet items
    if stripped.startswith(('- ', '* ', '• ')):
        para.style = doc.styles['List Bullet']
        run = para.add_run(stripped[2:])
    elif re.match(r'^\d+\.\s+', stripped) and not re.match(r'^\d+\.\d+', stripped):
        para.style = doc.styles['List Number']
        run = para.add_run(re.sub(r'^\d+\.\s+', '', stripped))
    else:
        # Indent detection
        indent = len(line) - len(line.lstrip())
        if indent >= 4:
            para.paragraph_format.left_indent = Pt(indent * 3)
        run = para.add_run(stripped)

    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    i += 1

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
